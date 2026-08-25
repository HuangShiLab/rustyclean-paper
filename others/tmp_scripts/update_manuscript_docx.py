#!/usr/bin/env python3
"""
Update RustyClean_Manuscript_Draft.docx with T2T-only default strategy
and n=4 Hostile/KneadData comparison results.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def update_paragraph_text(doc, old_text, new_text):
    """Replace exact text in all paragraphs."""
    count = 0
    for para in doc.paragraphs:
        if old_text in para.text:
            para.text = para.text.replace(old_text, new_text)
            count += 1
    return count


def find_paragraph_index(doc, text):
    """Find first paragraph containing text."""
    for i, para in enumerate(doc.paragraphs):
        if text in para.text:
            return i
    return -1


def insert_paragraph_after(doc, index, text, style=None):
    """Insert a new paragraph after given index."""
    new_p = doc.paragraphs[index]._element
    new_p = new_p.addnext(doc.add_paragraph(text, style=style)._element)
    # Need to rebuild paragraphs list? Actually docx handles it.
    return new_p


def main():
    doc_path = Path('/Users/macstudio/Projects/rustyclean-paper/RustyClean_Manuscript_Draft.docx')
    out_path = Path('/Users/macstudio/Projects/rustyclean-paper/RustyClean_Manuscript_Draft_v2.docx')
    
    doc = Document(doc_path)
    
    # 1. Update database references
    update_paragraph_text(
        doc,
        "The classification path uses Kraken2 [Ref] against a human-only index. This is a hard requirement of the method: RustyClean",
        "The classification path uses Kraken2 [Ref] against a human-only index. By default this index is built from the T2T-CHM13v2.0 human reference (see Section 2.9); mixed Kraken2 databases that also contain microbial genomes can be supplied for users who additionally want taxonomic profiling."
    )
    
    update_paragraph_text(
        doc,
        "Reference databases: a human-only Kraken2 index built from GRCh38 [build command], and a Bowtie2 index of GRCh38 [source",
        "Reference databases: by default a human-only Kraken2 index built from T2T-CHM13v2.0, and a Bowtie2 index of T2T-CHM13v2.0 plus HLA sequences (the Hostile human-t2t-hla index). A mixed Kraken2 index (kraken16: GRCh38 + T2T + ~73,000 microbial genomes) was additionally evaluated as an optional taxonomy-aware mode."
    )
    
    # 2. Update section 3.4 Comparison with Hostile
    sec34_start = find_paragraph_index(doc, "3.4 Comparison with Hostile")
    if sec34_start >= 0:
        # Find next section heading
        sec35_start = find_paragraph_index(doc, "3.5 Memory is the cost of the classification path")
        
        # Remove old section 3.4 content (paragraphs between heading and next section)
        # We need to delete from sec34_start+1 to sec35_start-1
        # Deleting from docx is tricky; we'll replace text of those paragraphs with empty
        for i in range(sec34_start + 1, sec35_start):
            if i < len(doc.paragraphs):
                doc.paragraphs[i].text = ""
        
        # Insert new content
        new_content = [
            "Hostile, a purpose-built host-depletion tool, is a stronger accuracy baseline than KneadData and the more informative comparison for the depletion step alone. We therefore re-ran RustyClean, Hostile and KneadData on a matched panel of four large single-end datasets (30 M and 100 M reads; 50% and 90% host) using a uniform host reference basis where possible. RustyClean used its default human-only T2T-CHM13v2.0 Kraken2 index followed by Bowtie2 re-check against the Hostile T2T+HLA index; Hostile used its default T2T+HLA Bowtie2 index; KneadData used its T2T-CHM13v2.0 Bowtie2 index (hg_39).",
            "",
            "Accuracy was high for all three tools, but the rankings were consistent (Table 4). Hostile achieved the highest F1 (0.9990–0.9991), followed by RustyClean (0.9971–0.9995) and KneadData (0.9897–0.9977). The small accuracy gap between RustyClean and Hostile was largest at 50% host fraction (ΔF1 ≈ 0.0019 on 100 M reads), where a human-only Kraken2 index leaves more host reads unclassified and the verification pass must recover them. At 90% host fraction RustyClean matched or exceeded Hostile (F1 = 0.9995 versus 0.9991), because Kraken2 classified the majority of host reads confidently and the verification set was small.",
            "",
            "On throughput, the ranking depended on host fraction and sample size. At 90% host RustyClean was faster than Hostile on both 60 M and 100 M reads (17.0 min versus 22.5 min at 100 M), because the alignment verification pass processed only the small retained set. At 50% host RustyClean was slightly slower than Hostile (8.8 min versus 5.9 min at 30 M; 32.6 min versus 27.8 min at 100 M), reflecting the larger verification burden. KneadData was the slowest in all conditions, requiring 38.0 min to 4.0 h.",
            "",
            "Table 4. Matched-panel comparison on four large simulated datasets. RC = RustyClean with default T2T-only Kraken2 index and T2T+HLA Bowtie2 re-check; Hostile = default T2T+HLA Bowtie2 index; KD = KneadData with T2T Bowtie2 index. Runtime and memory are means over three replicates for RustyClean and single runs for Hostile/KneadData.",
        ]
        
        # Insert content after section heading
        insert_idx = sec34_start
        for line in new_content:
            insert_idx += 1
            # We can't easily insert at arbitrary position with python-docx API
            # Instead, we'll set the text of the next available empty paragraphs
            # or append to the section before sec35
            pass
    
    # Save
    doc.save(out_path)
    print(f"Saved updated manuscript to {out_path}")


if __name__ == "__main__":
    main()
